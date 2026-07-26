"""
Professional Word document generation for patient invoices using python-docx.
"""
import io
import os
from decimal import Decimal

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from core.constants import CASHIER_NAME, DOCTOR_NAME


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_text(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def _add_table_row(table, cells_data, bold=False, header=False, bg_color=None):
    row = table.add_row()
    for i, (text, align) in enumerate(cells_data):
        cell = row.cells[i]
        _set_cell_text(cell, text, bold=bold, size=7.5,
                       align=align,
                       color=RGBColor(0xFF, 0xFF, 0xFF) if header else None)
        if bg_color:
            _set_cell_shading(cell, bg_color)
    return row


def _money(val):
    if val is None:
        return "—"
    return f"KSh {val:,.2f}"


def generate_invoice_docx(data):
    """Generate a professional A4 Word document invoice. Returns a BytesIO buffer."""
    buffer = io.BytesIO()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    hospital = data["hospital"]
    invoice = data["invoice"]
    patient = data["patient"]
    sm = data["summary"]

    # ── Hospital Header ──
    if hospital.logo and os.path.exists(hospital.logo.path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(hospital.logo.path, width=Inches(0.8))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hospital.name)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x21)

    if hospital.address:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hospital.address)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    contact = []
    if hospital.telephone:
        contact.append(f"Tel: {hospital.telephone}")
    if hospital.email:
        contact.append(f"Email: {hospital.email}")
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PATIENT INVOICE")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xDE, 0xE2, 0xE6)

    # ── Invoice & Patient Info (2-column table) ──
    def _add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    def _add_kv_table(rows_data):
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        for label, value in rows_data:
            row = t.add_row()
            _set_cell_text(row.cells[0], label, bold=True, size=8)
            _set_cell_text(row.cells[1], value, size=8)
            _set_cell_shading(row.cells[0], "F8F9FA")
        for row in t.rows:
            for cell in row.cells:
                cell.width = Inches(3.2)
        return t

    _add_section_heading("INVOICE & PATIENT INFORMATION")

    # Two-column layout
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    inv_data = [
        ("Invoice Number", invoice.invoice_number),
        ("Visit Number", invoice.visit.visit_number),
        ("Invoice Date", invoice.created_at.strftime("%d %b %Y %H:%M") if invoice.created_at else "—"),
        ("Cashier", CASHIER_NAME),
        ("Payment Status", invoice.get_status_display()),
    ]
    pat_data = [
        ("Patient Name", patient.full_name),
        ("Patient Number", patient.patient_number),
        ("Gender", patient.get_gender_display()),
        ("Age", f"{patient.age} years"),
        ("Phone", patient.phone),
        ("National ID", patient.national_id or "—"),
        ("Category", str(patient.patient_category) if patient.patient_category else "—"),
        ("Payment Type", patient.get_payment_type_display()),
    ]

    max_len = max(len(inv_data), len(pat_data))
    for i in range(max_len):
        row = t.add_row()
        if i < len(inv_data):
            _set_cell_text(row.cells[0], f"{inv_data[i][0]}: {inv_data[i][1]}", size=8)
            _set_cell_shading(row.cells[0], "F8F9FA")
        if i < len(pat_data):
            _set_cell_text(row.cells[1], f"{pat_data[i][0]}: {pat_data[i][1]}", size=8)

    # ── Visit Information ──
    _add_section_heading("VISIT INFORMATION")
    visit_rows = [
        ("Registration Date", invoice.visit.visit_date.strftime("%d %b %Y") if invoice.visit.visit_date else "—"),
        ("Consultation Date", data["consultation"].created_at.strftime("%d %b %Y") if data["consultation"] else "—"),
        ("Visit Type", data["visit_type"]),
        ("Attending Doctor", DOCTOR_NAME),
    ]
    if data["admission"]:
        adm = data["admission"]
        visit_rows.extend([
            ("Admission Date", adm.admission_date.strftime("%d %b %Y") if adm.admission_date else "—"),
            ("Discharge Date", adm.discharge_date.strftime("%d %b %Y") if adm.discharge_date else "Currently Admitted"),
            ("Ward", f"{adm.ward.name} ({adm.ward.get_ward_type_display()})"),
            ("Room / Bed", f"{adm.room.room_number} / {adm.bed.bed_number}"),
        ])
    _add_kv_table(visit_rows)

    # ── Diagnosis ──
    if data["primary_diagnosis"] or data["secondary_diagnosis"]:
        _add_section_heading("DIAGNOSIS")
        diag_rows = [("Primary Diagnosis", data["primary_diagnosis"])]
        if data["secondary_diagnosis"]:
            diag_rows.append(("Secondary Diagnosis", data["secondary_diagnosis"]))
        _add_kv_table(diag_rows)

    # ── Generic section table helper ──
    def _add_data_table(title, headers, rows, col_widths=None):
        if not rows:
            return
        _add_section_heading(title)
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, h in enumerate(headers):
            _set_cell_text(t.rows[0].cells[i], h, bold=True, size=7.5,
                           color=RGBColor(0xFF, 0xFF, 0xFF),
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_shading(t.rows[0].cells[i], "0D6EFD")

        for row_data in rows:
            row = t.add_row()
            for i, val in enumerate(row_data):
                _set_cell_text(row.cells[i], str(val), size=7.5)
        return t

    # ── Services ──
    if data["services"]:
        rows = [[s["date"], s["department"], s["service"], str(s["quantity"]),
                 _money(s["unit_price"]), _money(s["amount"])] for s in data["services"]]
        _add_data_table("SERVICES PROVIDED",
                        ["Date", "Department", "Service", "Qty", "Unit Price", "Amount"],
                        rows)

    # ── Lab Tests ──
    if data["lab_tests"]:
        rows = [[l["test_name"], str(l["quantity"]), _money(l["unit_price"]),
                 _money(l["amount"]), l["status"]] for l in data["lab_tests"]]
        _add_data_table("LABORATORY TESTS",
                        ["Test Name", "Qty", "Unit Price", "Amount", "Status"],
                        rows)

    # ── Radiology ──
    if data["radiology"]:
        rows = [[r["service"], str(r["quantity"]), _money(r["unit_price"]),
                 _money(r["amount"])] for r in data["radiology"]]
        _add_data_table("RADIOLOGY / IMAGING",
                        ["Service", "Qty", "Unit Price", "Amount"],
                        rows)

    # ── Medications ──
    if data["medications"]:
        rows = [[m["medicine_name"], m.get("dosage", ""), str(m["quantity"]),
                 _money(m["unit_price"]), _money(m["amount"])] for m in data["medications"]]
        _add_data_table("MEDICATIONS",
                        ["Medicine", "Dosage", "Qty", "Unit Price", "Amount"],
                        rows)

    # ── Ward Information ──
    if data["ward"]:
        w = data["ward"]
        _add_section_heading("WARD INFORMATION")
        ward_rows = [
            ("Ward Type", w["ward_type"]),
            ("Room Number", w["room_number"]),
            ("Bed Number", w["bed_number"]),
            ("Admission Date", w["admission_date"]),
            ("Discharge Date", w["discharge_date"]),
            ("Length of Stay", f"{w['length_of_stay']} night(s)"),
            ("Rate Per Night", _money(w["ward_rate"])),
            ("Total Ward Charges", _money(w["total_charges"])),
        ]
        _add_kv_table(ward_rows)

    # ── Nursing Services ──
    if data["nursing"]:
        rows = [[n["service"], str(n["quantity"]), _money(n["unit_price"]),
                 _money(n["amount"])] for n in data["nursing"]]
        _add_data_table("NURSING SERVICES",
                        ["Service", "Qty", "Unit Price", "Amount"],
                        rows)

    # ── Other Charges ──
    if data["other_charges"]:
        rows = [[o["service"], str(o["quantity"]), _money(o["unit_price"]),
                 _money(o["amount"])] for o in data["other_charges"]]
        _add_data_table("OTHER CHARGES",
                        ["Service", "Qty", "Unit Price", "Amount"],
                        rows)

    # ── Bill Summary ──
    _add_section_heading("BILL SUMMARY")
    sum_items = []
    for label, val in [
        ("Registration Fee", sm["registration_fee"]),
        ("Consultation Fee", sm["consultation_fee"]),
        ("Laboratory Charges", sm["lab_charges"]),
        ("Radiology Charges", sm["radiology_charges"]),
        ("Medication Charges", sm["medication_charges"]),
        ("Ward Charges", sm["ward_charges"]),
        ("Nursing Charges", sm["nursing_charges"]),
        ("Procedure Charges", sm["procedure_charges"]),
        ("Other Charges", sm["other_charges"]),
    ]:
        if val and val > 0:
            sum_items.append((label, _money(val)))

    sum_items.append(("Total Hospital Bill", _money(sm["grand_total"])))
    sum_items.append(("Amount Paid", _money(sm["amount_paid"])))
    sum_items.append(("Outstanding Before Rebate", _money(sm["outstanding_before_rebate"])))
    if sm["nhif_bed_rebate"] > 0:
        sum_items.append((f"NHIF/SHA Bed Rebate ({sm['nights_stayed']} nights @ {sm['rebate_per_night']}/night)", f"- {_money(sm['nhif_bed_rebate'])}"))
    sum_items.append(("Final Balance", _money(sm["balance"])))

    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, val in sum_items:
        row = t.add_row()
        _set_cell_text(row.cells[0], label, bold=True, size=8)
        align = WD_ALIGN_PARAGRAPH.RIGHT
        _set_cell_text(row.cells[1], val, size=8, bold=True, align=align)
        if label == "Total Hospital Bill":
            _set_cell_shading(row.cells[0], "E7F1FF")
            _set_cell_shading(row.cells[1], "E7F1FF")
        elif label == "NHIF/SHA Bed Rebate" or label.startswith("NHIF/SHA Bed Rebate"):
            _set_cell_shading(row.cells[0], "F0FFF4")
            _set_cell_shading(row.cells[1], "F0FFF4")
        elif label == "Final Balance":
            bg = "FFF5F5" if sm["balance"] > 0 else "F0FFF4"
            _set_cell_shading(row.cells[0], bg)
            _set_cell_shading(row.cells[1], bg)

    # ── Payment History ──
    if data["payments"]:
        rows = [[p["date"], p["receipt_number"], p["payment_method"],
                 _money(p["amount"]), p["cashier"]] for p in data["payments"]]
        _add_data_table("PAYMENT HISTORY",
                        ["Date", "Receipt #", "Method", "Amount", "Cashier"],
                        rows)

    # ── Footer ──
    doc.add_paragraph()
    for line in [
        "Thank you for choosing TENOCARE HOSPITAL.",
        "Please retain this invoice for future reference.",
        "This invoice was generated electronically by the TENOCARE HOSPITAL Information Management System.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(7.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x8B, 0x91, 0x9A)

    doc.save(buffer)
    buffer.seek(0)
    return buffer
